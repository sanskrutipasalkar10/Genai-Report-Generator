import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_complex_report(filename):
    doc = Document()
    
    # --- TITLE PAGE ---
    title = doc.add_heading('Global Renewable Energy Transition Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n')
    doc.add_paragraph('Date: January 15, 2025')
    doc.add_paragraph('Prepared by: Strategic Energy Insights Group')
    doc.add_paragraph('Classification: INTERNAL REVIEW')
    doc.add_page_break()

    # --- PAGE 1: EXECUTIVE SUMMARY ---
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        "The global shift towards renewable energy sources has accelerated significantly in the last fiscal year. "
        "Driven by policy mandates in the EU and market incentives in North America, solar and wind capacities "
        "have reached historic highs. However, grid modernization remains a critical bottleneck. "
        "This report analyzes the capacity additions, financial investments, and carbon reduction metrics "
        "across major economic zones."
    )
    doc.add_paragraph(
        "Our analysis indicates a 15% year-over-year growth in solar PV installations, while offshore wind "
        "projects faced supply chain headwinds. The following sections detail the regional breakdown of these "
        "developments and the associated financial implications."
    )
    
    doc.add_heading('2. Regional Capacity Analysis', level=1)
    doc.add_paragraph(
        "The installation capacity varies significantly by region. China continues to lead in total gigawatts (GW) "
        "installed, while Europe is focusing heavily on offshore wind efficiency. The table below summarizes "
        "the installed capacity for Solar and Wind energy as of Q4 2024."
    )

    # TABLE 1: Regional Capacity
    table1_data = [
        ['Region', 'Solar_Capacity_GW', 'Wind_Capacity_GW', 'Hydro_Capacity_GW', 'Growth_YoY'],
        ['North America', '150.5', '120.2', '80.0', '12%'],
        ['Europe', '180.0', '160.5', '95.0', '14%'],
        ['Asia Pacific', '300.2', '250.0', '150.5', '18%'],
        ['Latin America', '40.5', '25.0', '60.0', '8%'],
        ['Middle East', '35.0', '5.0', '2.0', '22%']
    ]
    _add_table(doc, table1_data, "Table 1: Regional Renewable Capacity (Q4 2024)")
    
    doc.add_paragraph(
        "As observed, the Asia Pacific region dominates due to aggressive manufacturing and state subsidies. "
        "The Middle East, while starting from a lower base, shows the highest percentage growth (22%) due to new "
        "mega-projects in solar infrastructure."
    )
    doc.add_page_break()

    # --- PAGE 2: FINANCIAL OVERVIEW ---
    doc.add_heading('3. Financial Investment Landscape', level=1)
    doc.add_paragraph(
        "Investment in renewable technologies requires substantial upfront capital. The cost per megawatt-hour (MWh) "
        "has decreased for solar, but battery storage costs remain volatile due to lithium pricing. "
        "The table below breaks down the capital expenditure (CapEx) and operational expenditure (OpEx) "
        "for major utility-scale projects."
    )
    
    # TABLE 2: Financials
    table2_data = [
        ['Technology_Type', 'Avg_CapEx_Million_USD', 'Avg_OpEx_Annual_USD', 'LCOE_USD_per_MWh', 'ROI_Period_Years'],
        ['Solar PV Utility', '0.85', '12,000', '35.00', '7.5'],
        ['Wind Onshore', '1.20', '25,000', '40.00', '8.2'],
        ['Wind Offshore', '2.50', '80,000', '75.00', '11.0'],
        ['Geothermal', '3.00', '100,000', '60.00', '12.5'],
        ['Biomass', '1.80', '150,000', '85.00', '9.0']
    ]
    _add_table(doc, table2_data, "Table 2: Financial Metrics by Technology Type")

    doc.add_paragraph(
        "Solar PV Utility remains the most cost-effective option with an LCOE of $35.00/MWh. Offshore wind, "
        "despite its higher CapEx ($2.5M/MW), offers higher capacity factors, justifying the longer ROI period of 11 years."
    )
    doc.add_page_break()

    # --- PAGE 3: ENVIRONMENTAL IMPACT ---
    doc.add_heading('4. Carbon Reduction & Emissions', level=1)
    doc.add_paragraph(
        "The primary driver for the transition is the reduction of CO2 emissions. We have tracked the "
        "carbon offset achieved by replacing coal and natural gas plants with renewable alternatives. "
        "The data below highlights the carbon intensity and total offset achieved in 2024."
    )

    # TABLE 3: Emissions
    table3_data = [
        ['Sector', 'CO2_Offset_Million_Tonnes', 'Coal_Displacement_MWh', 'Grid_Intensity_gCO2_kWh', 'Compliance_Rate'],
        ['Industrial', '500.5', '850,000', '120', '92%'],
        ['Residential', '200.2', '300,000', '95', '85%'],
        ['Transport', '150.0', '120,000', '110', '78%'],
        ['Commercial', '320.8', '550,000', '105', '95%']
    ]
    _add_table(doc, table3_data, "Table 3: Sectoral Carbon Reduction Performance")

    doc.add_paragraph(
        "The Industrial sector contributed the largest offset (500.5 Million Tonnes), driven by the electrification "
        "of heavy machinery. The Transport sector lags behind with a 78% compliance rate, indicating a need for "
        "faster EV adoption policies."
    )
    doc.add_page_break()

    # --- PAGE 4: FUTURE OUTLOOK ---
    doc.add_heading('5. Strategic Projections (2025-2030)', level=1)
    doc.add_paragraph(
        "Looking ahead, the integration of AI in grid management and the maturation of green hydrogen technologies "
        "will redefine the market. Our forecast model predicts the following growth trajectory for installed base "
        "across key markets."
    )

    # TABLE 4: Projections
    table4_data = [
        ['Market', 'Current_Base_GW', 'Proj_2026_GW', 'Proj_2028_GW', 'Proj_2030_GW', 'CAGR_Percent'],
        ['USA', '250', '290', '380', '500', '12.5%'],
        ['China', '400', '520', '750', '1100', '18.2%'],
        ['EU-27', '300', '360', '480', '650', '14.0%'],
        ['India', '120', '160', '240', '350', '19.5%']
    ]
    _add_table(doc, table4_data, "Table 4: 5-Year Growth Projections")

    doc.add_paragraph(
        "India and China are expected to exhibit the highest CAGR at 19.5% and 18.2% respectively. "
        "The US market growth is steady but contingent on federal tax credit extensions."
    )
    
    doc.add_heading('6. Conclusion', level=1)
    doc.add_paragraph(
        "The data confirms that while financial barriers exist for offshore wind and geothermal, solar PV "
        "continues to democratize energy access. Immediate strategic focus should be on grid modernization "
        "to handle the intermittency of these expanded renewable capacities."
    )

    # Ensure directory exists
    os.makedirs(os.path.dirname(filename), exist_ok=True)
    doc.save(filename)
    print(f"✅ Generated Complex DOCX Report: {filename}")

def _add_table(doc, data, caption):
    doc.add_paragraph(caption, style='Caption')
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = 'Table Grid'
    
    for i, row in enumerate(data):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.text = str(val)
            # Bold header
            if i == 0:
                run = cell.paragraphs[0].runs[0]
                run.font.bold = True

if __name__ == "__main__":
    create_complex_report("data/raw/complex_energy_report.docx")