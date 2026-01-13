from docx import Document
from docx.shared import Pt

def create_messy_docx(filename):
    doc = Document()
    
    # 1. Add Random Text
    doc.add_heading('Q3 Financial Overview', 0)
    doc.add_paragraph('The following table contains the un-audited figures for the quarter.')
    doc.add_paragraph('Note: Data is subject to change.')

    # 2. Add a "Messy" Table
    # Structure: 
    # Row 0: "INTERNAL ONLY" (Garbage Metadata)
    # Row 1: "Category", "Value", "Value" (Duplicate Header)
    # Row 2: "Revenue", "100,000", "500"
    # Row 3: "Cost", "60,000", "200"
    
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    # --- Row 0: Garbage Metadata ---
    # Merging cells manually by putting text in first cell
    row0 = table.rows[0].cells
    row0[0].text = "INTERNAL USE ONLY - DO NOT DISTRIBUTE"
    # (In real docx, this might be a merged cell, but text in col 0 simulates the sanitizer challenge)

    # --- Row 1: Duplicate Headers ---
    row1 = table.rows[1].cells
    row1[0].text = "Category"
    row1[1].text = "Value"       # Duplicate 1
    row1[2].text = "Value"       # Duplicate 2 (Sanitizer should rename to Value_1)

    # --- Row 2: Data ---
    row2 = table.rows[2].cells
    row2[0].text = "Revenue"
    row2[1].text = "100,000"     # String number with comma
    row2[2].text = "500"

    # --- Row 3: Data ---
    row3 = table.rows[3].cells
    row3[0].text = "Cost"
    row3[1].text = "60,000"
    row3[2].text = "200"

    doc.save(filename)
    print(f"✅ Generated messy DOCX: {filename}")

if __name__ == "__main__":
    create_messy_docx("data/raw/test_messy.docx")