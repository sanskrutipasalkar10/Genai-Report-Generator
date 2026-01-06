import pandas as pd
import os
import docx  # Requires: pip install python-docx
from typing import Tuple, Dict, Any
from src.engine.agents.inspector import inspect_file

# Import your existing PDF parser (assuming it handles text+tables)
try:
    from src.rag.parsers.pdf_parser import parse_hybrid_pdf
except ImportError:
    parse_hybrid_pdf = None  # Fallback if parser isn't set up

def load_file(file_path: str) -> Tuple[str, Dict[str, pd.DataFrame], Dict[str, Any]]:
    """
    Universal File Loader.
    Supports: XLSX, XLS, CSV, PDF, DOCX.
    Maximizes data reading (500 rows) for Cloud Models.
    """
    ext = os.path.splitext(file_path)[1].lower()
    filename = os.path.basename(file_path)
    print(f"📂 Processing: {filename}")

    # 🟢 CONFIG: Cloud Model Context Window
    # The 480B model can easily handle 500 rows of context.
    INSPECT_ROWS = 500 
    
    raw_preview = ""
    tables = {}
    full_text = ""

    try:
        # --- 1. EXCEL FILES ---
        if ext in [".xlsx", ".xls"]:
            # Load Dataframes (Dictionary of sheets)
            tables = pd.read_excel(file_path, sheet_name=None)
            
            # Generate Preview from the largest sheet (most likely containing data)
            if tables:
                largest_sheet = max(tables, key=lambda k: tables[k].size)
                df_preview = tables[largest_sheet].head(INSPECT_ROWS)
                raw_preview = f"Sheet '{largest_sheet}' Preview (First {INSPECT_ROWS} Rows):\n{df_preview.to_string()}"
            
        # --- 2. CSV FILES ---
        elif ext == ".csv":
            # Load Data
            df = pd.read_csv(file_path)
            tables = {"CSV_Data": df}
            
            # Generate Preview
            df_preview = df.head(INSPECT_ROWS)
            raw_preview = f"CSV Preview (First {INSPECT_ROWS} Rows):\n{df_preview.to_string()}"

        # --- 3. PDF FILES ---
        elif ext == ".pdf":
            if parse_hybrid_pdf:
                full_text, pdf_tables = parse_hybrid_pdf(file_path)
                # Convert list of DFs to a dictionary
                tables = {f"Table_{i}": t for i, t in enumerate(pdf_tables)}
                
                # Preview: Text + First Table
                table_preview = pdf_tables[0].head(50).to_string() if pdf_tables else "No Tables Found"
                raw_preview = f"TEXT PREVIEW:\n{full_text[:1000]}\n\nTABLE PREVIEW:\n{table_preview}"
            else:
                raw_preview = "Error: PDF Parser not found."

        # --- 4. WORD DOCUMENTS (.docx) ---
        elif ext == ".docx":
            doc = docx.Document(file_path)
            
            # Extract Text
            full_text = "\n".join([para.text for para in doc.paragraphs])
            
            # Extract Tables
            for i, table in enumerate(doc.tables):
                data = [[cell.text for cell in row.cells] for row in table.rows]
                if data:
                    # Assume first row is header
                    df = pd.DataFrame(data[1:], columns=data[0])
                    tables[f"Word_Table_{i}"] = df
            
            # Preview
            table_preview = list(tables.values())[0].head(50).to_string() if tables else "No Tables"
            raw_preview = f"TEXT PREVIEW:\n{full_text[:1000]}\n\nTABLE PREVIEW:\n{table_preview}"

        else:
            print(f"❌ Unsupported file type: {ext}")
            return "", {}, {}

    except Exception as e:
        print(f"❌ Read Error: {e}")
        return "", {}, {}

    # --- 5. INTELLIGENT INSPECTION ---
    # Send the massive preview to the Inspector Agent
    config = inspect_file(raw_preview, filename)

    # --- 6. GENERALIZED CLEANING ---
    # Prepare tables for the Analyst based on Inspector's findings
    cleaned_tables = {}
    for name, df in tables.items():
        if isinstance(df, pd.DataFrame) and not df.empty:
            
            # Auto-Cleanup: Drop completely empty rows/cols
            df.dropna(how='all', axis=0, inplace=True)
            df.dropna(how='all', axis=1, inplace=True)
            
            # Report Specific Logic
            if config.get("report_type") == "FINANCIAL":
                # Ensure we don't accidentally overwrite useful headers
                # Only rename if the first column looks like a label index
                if "line_item" not in [str(c).lower() for c in df.columns]:
                     # Heuristic: If col 0 is object/string, treat as Labels
                     if df.dtypes.iloc[0] == 'object':
                        df.rename(columns={df.columns[0]: "Line_Item"}, inplace=True)

            cleaned_tables[name] = df

    return full_text, cleaned_tables, config