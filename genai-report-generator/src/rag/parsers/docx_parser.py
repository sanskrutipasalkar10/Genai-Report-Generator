import pandas as pd
import docx
from typing import Tuple, Dict

def parse_docx(file_path: str) -> Tuple[str, Dict[str, pd.DataFrame]]:
    """
    Parses a DOCX file to extract text and VALID data tables.
    Returns raw DataFrames (no headers set) for the Sanitizer to clean.
    """
    print(f"   📄 Parsing DOCX: {file_path}")
    
    try:
        doc = docx.Document(file_path)
    except Exception as e:
        print(f"   ❌ DOCX Load Error: {e}")
        return "", {}

    full_text = []
    tables = {}

    # 1. Extract Text
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)

    # 2. Extract & Filter Tables
    for i, table in enumerate(doc.tables):
        try:
            # EXTRACT: Get all rows as list of lists
            data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            
            if not data:
                continue

            # --- 🟢 HEURISTIC FILTERING ---
            
            # Rule 1: Minimum Size (Must have at least 2 rows and 2 columns)
            if len(data) < 2 or len(data[0]) < 2:
                continue
            
            # Rule 2: Check for "Ghost" tables (mostly empty)
            flat_cells = [item for sublist in data for item in sublist]
            if not flat_cells: continue
            
            empty_count = flat_cells.count("")
            if empty_count / len(flat_cells) > 0.9: # Skip if >90% empty
                continue
            
            # Create DataFrame
            # 🟢 CRITICAL CHANGE: header=None. 
            # We let DataSanitizer decide which row is the header later.
            df = pd.DataFrame(data)
            
            # Rule 3: Content check - Financial tables usually have numbers
            # Check if at least one cell contains a digit
            has_numbers = df.astype(str).apply(lambda x: x.str.contains(r'\d', na=False)).any().any()
            
            if has_numbers:
                tables[f"Table_{i+1}"] = df
                
        except Exception as e:
            print(f"      ⚠️ Warning processing DOCX table {i}: {e}")

    print(f"      ✅ Extracted {len(tables)} tables from DOCX.")
    return "\n".join(full_text), tables